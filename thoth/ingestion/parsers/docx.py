"""Word document parser.

This module provides parsing for Word documents using python-docx.
"""

import logging
from pathlib import Path
from typing import Any

from thoth.ingestion.parsers.base import DocumentParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    """Parser for Word documents using python-docx.

    Supports:
    - Word documents (.docx)
    - Paragraph text extraction
    - Basic metadata extraction (title, author)

    Note:
        Only supports .docx format (Office Open XML).
        Legacy .doc files are not supported.
    """

    @property
    def supported_extensions(self) -> list[str]:
        """Return supported Word document extensions."""
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a Word document.

        Args:
            file_path: Path to the Word document

        Returns:
            ParsedDocument with extracted text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ImportError: If python-docx is not installed
        """
        if not file_path.exists():
            msg = f"File not found: {file_path}"
            raise FileNotFoundError(msg)

        try:
            from docx import Document  # noqa: PLC0415
        except ImportError as e:
            msg = "python-docx is required for Word document parsing. Install with: pip install python-docx"
            raise ImportError(msg) from e

        doc = Document(str(file_path))
        return self._extract_document(doc, str(file_path))

    def parse_content(self, content: bytes, source_path: str) -> ParsedDocument:
        """Parse Word document content from bytes.

        Args:
            content: Raw document content as bytes
            source_path: Original source path for metadata

        Returns:
            ParsedDocument with extracted text and metadata
        """
        try:
            from docx import Document  # noqa: PLC0415
        except ImportError as e:
            msg = "python-docx is required for Word document parsing. Install with: pip install python-docx"
            raise ImportError(msg) from e

        import io  # noqa: PLC0415

        doc = Document(io.BytesIO(content))
        return self._extract_document(doc, source_path)

    def _extract_document(self, doc: Any, source_path: str) -> ParsedDocument:
        """Extract text and metadata from a Document object.

        Args:
            doc: python-docx Document object
            source_path: Original source path for metadata

        Returns:
            ParsedDocument with extracted content
        """
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        # Extract metadata from core properties
        core_props = doc.core_properties
        metadata = {
            "source_path": source_path,
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "paragraph_count": len(paragraphs),
        }

        # Remove empty metadata values
        metadata = {k: v for k, v in metadata.items() if v}

        return ParsedDocument(
            content="\n\n".join(paragraphs),
            metadata=metadata,
            source_path=source_path,
            format="docx",
        )
